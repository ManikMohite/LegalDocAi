# from flask import Flask, request, jsonify, render_template, session, send_file
# from flask_cors import CORS
# import google.generativeai as genai
# import PyPDF2
# import os
# from dotenv import load_dotenv
# from io import BytesIO
# import tempfile
# from docx import Document
# from docx.shared import Pt, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# import uuid
# import datetime
# from multiagent import *


# load_dotenv()

# api_key = os.getenv('GEMINI_API_KEY')

# app = Flask(__name__)

# app.secret_key = os.getenv('SECRET_KEY', 'default-secret-key')
# CORS(app,
#      supports_credentials=True,
#      resources={r"/*": {"origins": "http://localhost:3000"}}
# )

# general_context = ""
# doc_chat_context = ""

# # Configure Gemini API key
# genai.configure(api_key=api_key)

# # Store document text and files globally
# document_cache = {}
# pdf_cache = {}
# draft_cache = {}

# # ------------------ Category metrics ------------------
# CATEGORY_METRICS = {
#     'Legal Notice': [
#         'Severity Score', 'Violations & Broken Rules', 'Legal Consequences', 'Actionable Steps',
#         'Urgency Detection', 'Tone Analysis', 'Recommended Actions'
#     ],
#     'Ownership Documents': [
#         'Ownership Rights & Obligations', 'Transfer, Leasing, Sale, Mortgaging Clauses',
#         'Financial Liabilities', 'Terms & Conditions', 'Important Dates', 'Document Validity', 'Summary Type'
#     ],
#     'Contracts & Agreements': [
#         'Parties Involved & Roles', 'Terms & Conditions', 'Termination Clauses', 'Penalties for Breach',
#         'Severity Score', 'Obligations & Rights', 'Actionable Steps'
#     ],
#     'Financial Documents': [
#         'Financial Obligations', 'Coverage Details', 'Deadlines & Payment Schedules',
#         'Legal Implications', 'Severity Score', 'Urgency Detection', 'Risk Analysis'
#     ],
#     'Terms & Conditions / Privacy Policies': [
#         'User Rights & Restrictions', 'Data Usage & Privacy Clauses', 'Liability Clauses',
#         'Termination & Suspension Rules', 'Severity Score', 'Personal Implications', 'Suggested Actions'
#     ],
#     'Intellectual Property Documents': [
#         'Ownership & Usage Rights', 'Infringement Clauses', 'Exclusivity & Licensing Terms',
#         'Penalties for Violation', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
#     ],
#     'Criminal Offense Documents': [
#         'Charges Filed', 'Potential Penalties', 'Required Actions', 'Severity Score',
#         'Urgency Detection', 'Tone Analysis', 'Suggested Actions'
#     ],
#     'Regulatory Compliance Documents': [
#         'Compliance Requirements', 'Penalties for Non-Compliance', 'Renewal Deadlines & Conditions',
#         'Guidelines for Rectification', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
#     ],
#     'Employment Documents': [
#         'Terms of Employment', 'Termination Conditions', 'Confidentiality Clauses',
#         'Breach Consequences', 'Severity Score', 'Urgency Detection', 'Suggested Actions'
#     ],
#     'Court Judgments & Legal Precedents': [
#         'Summary of Judgment', 'Legal Basis', 'Potential Consequences',
#         'Severity Score', 'Urgency Detection', 'Recommended Actions'
#     ]
# }

# # ------------------ Draft templates ------------------
# DRAFT_TEMPLATES = {
#     'Legal Notice Response': {
#         'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
#         'header_format': {'font': 'Times New Roman', 'size': 12, 'bold': True, 'align': 'center'},
#         'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
#         'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
#         'date_format': '%B %d, %Y',
#         'includes_header': True,
#         'includes_date': True,
#         'includes_signature': True
#     },
#     'Contract Response': {
#         'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
#         'header_format': {'font': 'Arial', 'size': 12, 'bold': True, 'align': 'center'},
#         'body_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
#         'signature_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
#         'date_format': '%d/%m/%Y',
#         'includes_header': True,
#         'includes_date': True,
#         'includes_signature': True
#     },
#     'General Letter': {
#         'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
#         'header_format': {'font': 'Calibri', 'size': 12, 'bold': True, 'align': 'left'},
#         'body_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
#         'signature_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
#         'date_format': '%B %d, %Y',
#         'includes_header': True,
#         'includes_date': True,
#         'includes_signature': True
#     },
#     'Legal Memo': {
#         'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
#         'header_format': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'center'},
#         'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
#         'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
#         'date_format': '%B %d, %Y',
#         'includes_header': True,
#         'includes_date': True,
#         'includes_signature': False
#     }
# }

# CATEGORY_TO_TEMPLATE = {
#     'Legal Notice': 'Legal Notice Response',
#     'Contracts & Agreements': 'Contract Response',
#     'Ownership Documents': 'Legal Memo',
#     'Financial Documents': 'Legal Memo',
#     'Terms & Conditions / Privacy Policies': 'Legal Memo',
#     'Intellectual Property Documents': 'Legal Memo',
#     'Criminal Offense Documents': 'Legal Notice Response',
#     'Regulatory Compliance Documents': 'Legal Memo',
#     'Employment Documents': 'Legal Memo',
#     'Court Judgments & Legal Precedents': 'Legal Memo',
#     'default': 'General Letter'
# }

# # ------------------ Routes ------------------
# @app.route('/')
# def home():
#     return jsonify({"status": "Backend running"}), 200


# # @app.route('/general_chat.html')
# # def general_chat():
# #     return render_template('general_chat.html')

# def extract_text_from_pdf(pdf_file):
#     try:
#         pdf_content = BytesIO(pdf_file.read())
#         pdf_file.seek(0)
#         reader = PyPDF2.PdfReader(pdf_content)
#         text = ''
#         for page_num in range(len(reader.pages)):
#             text += reader.pages[page_num].extract_text()
#         return text
#     except Exception as e:
#         app.logger.error(f"PDF extraction error: {str(e)}")
#         return str(e)

# # ------------------ Document Classification ------------------
# @app.route('/classify', methods=['POST'])
# def classify_document():
#     if 'document' not in request.files:
#         return jsonify({'error': 'No PDF file uploaded'}), 400

#     pdf_file = request.files['document']
#     session_id = session.get('session_id', os.urandom(16).hex())
#     pdf_content = pdf_file.read()
#     pdf_cache[session_id] = pdf_content
#     pdf_file.seek(0)

#     document_text = extract_text_from_pdf(pdf_file)
#     if not document_text:
#         return jsonify({'error': 'Failed to extract text from PDF'}), 400

#     document_cache[session_id] = document_text

#     try:
#         model = genai.GenerativeModel("gemini-2.0-flash")
#         response = model.generate_content([
#             "You are a document classification agent. Classify the document into one of these categories: Legal Notice, Ownership Documents, Contracts & Agreements, Financial Documents, Terms & Conditions / Privacy Policies, Intellectual Property Documents, Criminal Offense Documents, Regulatory Compliance Documents, Employment Documents, Court Judgments & Legal Precedents.",
#             document_text[:3000]
#         ])
#         category = response.text.strip()
#         return jsonify({'category': category})
#     except Exception as e:
#         app.logger.error(f"Classification error: {str(e)}")
#         return jsonify({'error': str(e)}), 500

# # ------------------ Document Processing ------------------
# @app.route('/process', methods=['POST'])
# def process_document():
#     if 'document' not in request.files or 'category' not in request.form:
#         return jsonify({'error': 'Document file or category is missing'}), 400

#     pdf_file = request.files['document']
#     category = request.form['category']
#     document_text = extract_text_from_pdf(pdf_file)
#     if not document_text or not category:
#         return jsonify({'error': 'Document text or category is missing'}), 400

#     session_id = session.get('session_id', os.urandom(16).hex())
#     document_cache[session_id] = document_text

#     metrics = CATEGORY_METRICS.get(category, [])
#     metrics_prompt = ', '.join(metrics)
#     prompt = f"You are an expert summarizer for {category} documents. Extract the following relevant metrics: {metrics_prompt}. Format each metric as '**Metric Name**: Value' to make it bold and easily readable."

#     try:
#         model = genai.GenerativeModel("gemini-2.0-flash")
#         response = model.generate_content([
#             prompt,
#             document_text[:4000]
#         ])
#         summary = response.text.strip()
#         return jsonify({
#             'summary': summary,
#             'document_text': document_text[:200] + '...' if len(document_text) > 200 else document_text
#         })
#     except Exception as e:
#         app.logger.error(f"Processing error: {str(e)}")
#         return jsonify({'error': str(e)}), 500

# # ------------------ Chat ------------------


# @app.route('/chat', methods=['POST'])
# def chat():
#     # ensure session id exists
#     if 'session_id' not in session:
#         session['session_id'] = os.urandom(16).hex()

#     global doc_chat_context

#     # safer JSON parsing (won't raise 400)
#     data = request.get_json(silent=True)
#     if data is None:
#         # helpful debug logs
#         app.logger.debug("Bad request to /chat: no JSON. raw data: %s", request.data)
#         app.logger.debug("Headers: %s", dict(request.headers))
#         return jsonify({'error': 'Invalid or missing JSON in request body. Make sure you send Content-Type: application/json and a valid JSON body.'}), 400

#     user_message = data.get('message')
#     category = data.get('category')
#     detailed_analysis = data.get('detailed_analysis', False)
#     generate_draft_flag = data.get('generate_draft', False)
#     draft_instructions = data.get('draft_instructions', '')

#     if not user_message:
#         return jsonify({'error': 'Message is required'}), 400

#     # ensure session id is stored
#     session_id = session.get('session_id')
#     if not session_id:
#         session_id = os.urandom(16).hex()
#     session['session_id'] = session_id

#     document_text = document_cache.get(session_id, '')
#     if not document_text:
#         return jsonify({'error': 'No document found. Please process a document first.'}), 400

#     try:
#         document_context = document_text[:3000] if len(document_text) > 3000 else document_text
#         system_prompt = f"""You are a legal assistant specializing in {category} documents.
# Document context (truncated): {document_context}
# Chat context: {doc_chat_context}
# """
#         if detailed_analysis:
#             system_prompt += "Provide a detailed analysis with comprehensive explanations, legal references, and thorough examination of all relevant aspects."
#         else:
#             system_prompt += "Provide concise, clear answers focused on the most important points."
#         if generate_draft_flag:
#             draft_id = generate_document_draft(user_message, draft_instructions, category, document_context)
#             return jsonify({
#                 'response': "I've prepared a draft document based on your instructions. You can download it using the link below.",
#                 'draft_id': draft_id
#             })

#         system_prompt += "Provide helpful, accurate information based on this document. Use **bold** for important points."

#         model = genai.GenerativeModel("gemini-2.0-flash")
#         response = model.generate_content([
#             system_prompt,
#             user_message
#         ])
#         bot_response = response.text.strip()
#         doc_chat_context += f"\nUser: {user_message}\nBot: {bot_response}\n"
#         return jsonify({'response': bot_response})

#     except Exception as e:
#         app.logger.error(f"Chat error: {str(e)}", exc_info=True)
#         return jsonify({'error': str(e)}), 500

# # ------------------ General Chat ------------------
# @app.route('/general_chat', methods=['POST'])
# def general_chat_api():
#     global general_context
#     data = request.json
#     user_message = data.get('message')
#     detailed_analysis = data.get('detailed_analysis', False)

#     if not user_message:
#         return jsonify({'error': 'Message is required'}), 400

#     try:
#         if detailed_analysis:
#             response, reasoning = get_answer(user_message, general_context)
#             general_context += f"\n User: {user_message}\nSenior Lawyer: {response}\n"
#             return jsonify({'response': response, 'reasoning': reasoning})
#         else:
#             system_prompt = f"""You are a knowledgeable legal assistant providing general legal information.
# You are not a lawyer and should clarify that your responses do not constitute legal advice.
# Context: {general_context}
# Provide concise, clear answers with **bold** important points.
# """
#             model = genai.GenerativeModel("gemini-2.0-flash")
#             response = model.generate_content([
#                 system_prompt,
#                 user_message
#             ])
#             bot_response = response.text.strip()
#             return jsonify({'response': bot_response, 'reasoning': []})
#     except Exception as e:
#         app.logger.error(f"General chat error: {str(e)}")
#         return jsonify({'error': str(e)}), 500

# # ------------------ Draft Generation ------------------
# def generate_document_draft(message, instructions, category, document_context):
#     template_name = CATEGORY_TO_TEMPLATE.get(category, CATEGORY_TO_TEMPLATE['default'])
#     template = DRAFT_TEMPLATES[template_name]

#     prompt = f"""You are a professional legal document drafter. Create a formal response document based on the instructions:
# Instructions: {instructions}
# Category: {category}
# Context (truncated): {document_context[:1500]}
# Format: header, date, salutation, body, closing, signature as per professional standards.
# """

#     try:
#         model = genai.GenerativeModel("gemini-2.0-flash")
#         response = model.generate_content([prompt, message])
#         draft_content = response.text.strip()

#         doc = create_formatted_document(draft_content, template)
#         draft_id = str(uuid.uuid4())
#         temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
#         doc.save(temp_file.name)
#         draft_cache[draft_id] = {
#             'path': temp_file.name,
#             'filename': f"Legal_Draft_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
#         }
#         return draft_id
#     except Exception as e:
#         app.logger.error(f"Draft generation error: {str(e)}")
#         raise

# # ------------------ Document Formatting ------------------
# def create_formatted_document(content, template):
#     doc = Document()
#     for section in doc.sections:
#         section.top_margin = Inches(template['margins']['top'])
#         section.bottom_margin = Inches(template['margins']['bottom'])
#         section.left_margin = Inches(template['margins']['left'])
#         section.right_margin = Inches(template['margins']['right'])

#     lines = content.split('\n')
#     current_section = 'header'

#     for line in lines:
#         line = line.strip()
#         if not line:
#             doc.add_paragraph()
#             continue
#         p = doc.add_paragraph()
#         if current_section == 'header':
#             p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template['header_format']['align'] == 'center' else WD_ALIGN_PARAGRAPH.LEFT
#             run = p.add_run(line)
#             run.font.name = template['header_format']['font']
#             run.font.size = Pt(template['header_format']['size'])
#             run.font.bold = template['header_format']['bold']
#             current_section = 'body'
#         elif current_section == 'body':
#             p.alignment = WD_ALIGN_PARAGRAPH.LEFT
#             run = p.add_run(line)
#             run.font.name = template['body_format']['font']
#             run.font.size = Pt(template['body_format']['size'])
#         elif current_section == 'signature':
#             p.alignment = WD_ALIGN_PARAGRAPH.LEFT
#             run = p.add_run(line)
#             run.font.name = template['signature_format']['font']
#             run.font.size = Pt(template['signature_format']['size'])

#     return doc

# # ------------------ Download and View ------------------
# @app.route('/download-draft/<draft_id>', methods=['GET'])
# def download_draft(draft_id):
#     if draft_id not in draft_cache:
#         return jsonify({'error': 'Draft not found'}), 404
#     draft_info = draft_cache[draft_id]
#     try:
#         return send_file(
#             draft_info['path'],
#             as_attachment=True,
#             download_name=draft_info['filename'],
#             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#         )
#     except Exception as e:
#         app.logger.error(f"Draft download error: {str(e)}")
#         return jsonify({'error': 'Error downloading draft'}), 500

# @app.route('/view-document', methods=['GET'])
# def view_document():
#     session_id = session.get('session_id')
#     if not session_id or session_id not in pdf_cache:
#         return jsonify({'error': 'No document found'}), 404
#     temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
#     temp_file.write(pdf_cache[session_id])
#     temp_file.close()
#     return send_file(temp_file.name, mimetype='application/pdf', as_attachment=False)

# # ------------------ Run App ------------------
# if __name__ == '__main__':
#     app.run(debug=True)

from flask import Flask, request, jsonify, session, send_file
from flask_cors import CORS
import google.generativeai as genai
import PyPDF2
import os
from dotenv import load_dotenv
from io import BytesIO
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import datetime
import logging

# load environment
load_dotenv()
api_key = os.getenv('GEMINI_API_KEY')

# app init
app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'default-secret-key')

# CORS - allow frontend to send cookies
CORS(app,
     supports_credentials=True,
     resources={r"/*": {"origins": [
         "http://localhost:3000",
         "http://127.0.0.1:3000"
     ]}}
)


# basic logging
logging.basicConfig(level=logging.DEBUG)
app.logger.setLevel(logging.DEBUG)

# Configure Gemini API key (if present)
if api_key:
    try:
        genai.configure(api_key=api_key)
        app.logger.info("Gemini API configured")
    except Exception as e:
        app.logger.warning(f"Failed to configure Gemini API: {e}")
else:
    app.logger.warning("GEMINI_API_KEY not set; GenAI calls will likely fail in production")

# global caches
general_context = ""
doc_chat_context = ""
document_cache = {}
pdf_cache = {}
draft_cache = {}

# Category metrics and templates (unchanged)
CATEGORY_METRICS = {
    'Legal Notice': [
        'Severity Score', 'Violations & Broken Rules', 'Legal Consequences', 'Actionable Steps',
        'Urgency Detection', 'Tone Analysis', 'Recommended Actions'
    ],
    'Ownership Documents': [
        'Ownership Rights & Obligations', 'Transfer, Leasing, Sale, Mortgaging Clauses',
        'Financial Liabilities', 'Terms & Conditions', 'Important Dates', 'Document Validity', 'Summary Type'
    ],
    'Contracts & Agreements': [
        'Parties Involved & Roles', 'Terms & Conditions', 'Termination Clauses', 'Penalties for Breach',
        'Severity Score', 'Obligations & Rights', 'Actionable Steps'
    ],
    'Financial Documents': [
        'Financial Obligations', 'Coverage Details', 'Deadlines & Payment Schedules',
        'Legal Implications', 'Severity Score', 'Urgency Detection', 'Risk Analysis'
    ],
    'Terms & Conditions / Privacy Policies': [
        'User Rights & Restrictions', 'Data Usage & Privacy Clauses', 'Liability Clauses',
        'Termination & Suspension Rules', 'Severity Score', 'Personal Implications', 'Suggested Actions'
    ],
    'Intellectual Property Documents': [
        'Ownership & Usage Rights', 'Infringement Clauses', 'Exclusivity & Licensing Terms',
        'Penalties for Violation', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ],
    'Criminal Offense Documents': [
        'Charges Filed', 'Potential Penalties', 'Required Actions', 'Severity Score',
        'Urgency Detection', 'Tone Analysis', 'Suggested Actions'
    ],
    'Regulatory Compliance Documents': [
        'Compliance Requirements', 'Penalties for Non-Compliance', 'Renewal Deadlines & Conditions',
        'Guidelines for Rectification', 'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ],
    'Employment Documents': [
        'Terms of Employment', 'Termination Conditions', 'Confidentiality Clauses',
        'Breach Consequences', 'Severity Score', 'Urgency Detection', 'Suggested Actions'
    ],
    'Court Judgments & Legal Precedents': [
        'Summary of Judgment', 'Legal Basis', 'Potential Consequences',
        'Severity Score', 'Urgency Detection', 'Recommended Actions'
    ]
}

DRAFT_TEMPLATES = {
    'Legal Notice Response': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Times New Roman', 'size': 12, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'date_format': '%B %d, %Y',
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'Contract Response': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Arial', 'size': 12, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
        'signature_format': {'font': 'Arial', 'size': 11, 'align': 'left'},
        'date_format': '%d/%m/%Y',
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'General Letter': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Calibri', 'size': 12, 'bold': True, 'align': 'left'},
        'body_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
        'signature_format': {'font': 'Calibri', 'size': 11, 'align': 'left'},
        'date_format': '%B %d, %Y',
        'includes_header': True,
        'includes_date': True,
        'includes_signature': True
    },
    'Legal Memo': {
        'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.25, 'right': 1.25},
        'header_format': {'font': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'center'},
        'body_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'signature_format': {'font': 'Times New Roman', 'size': 12, 'align': 'left'},
        'date_format': '%B %d, %Y',
        'includes_header': True,
        'includes_date': True,
        'includes_signature': False
    }
}

CATEGORY_TO_TEMPLATE = {
    'Legal Notice': 'Legal Notice Response',
    'Contracts & Agreements': 'Contract Response',
    'Ownership Documents': 'Legal Memo',
    'Financial Documents': 'Legal Memo',
    'Terms & Conditions / Privacy Policies': 'Legal Memo',
    'Intellectual Property Documents': 'Legal Memo',
    'Criminal Offense Documents': 'Legal Notice Response',
    'Regulatory Compliance Documents': 'Legal Memo',
    'Employment Documents': 'Legal Memo',
    'Court Judgments & Legal Precedents': 'Legal Memo',
    'default': 'General Letter'
}

# ------------------ Helpers ------------------

def _ensure_session_id():
    sid = session.get('session_id')
    if not sid:
        sid = os.urandom(16).hex()
        session['session_id'] = sid
        app.logger.debug(f"New session created: {sid}")
    return sid


def extract_text_from_pdf_bytes(pdf_bytes):
    try:
        reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        text = ''
        for page in reader.pages:
            page_text = page.extract_text() or ''
            text += page_text
        return text
    except Exception as e:
        app.logger.error(f"PDF extraction error: {e}")
        return ''

# ------------------ Routes ------------------

@app.route('/')
def home():
    return jsonify({"status": "Backend running"}), 200


@app.route('/session', methods=['GET'])
def get_session():
    """Debug endpoint to see current session id (for local dev)."""
    sid = session.get('session_id')
    return jsonify({'session_id': sid}), 200


@app.route('/classify', methods=['POST'])
def classify_document():
    # ensure session
    sid = _ensure_session_id()

    if 'document' not in request.files:
        return jsonify({'error': 'No PDF file uploaded'}), 400

    pdf_file = request.files['document']
    pdf_bytes = pdf_file.read()
    pdf_cache[sid] = pdf_bytes

    document_text = extract_text_from_pdf_bytes(pdf_bytes)
    if not document_text:
        return jsonify({'error': 'Failed to extract text from PDF'}), 400

    document_cache[sid] = document_text

    try:
        if api_key:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content([
                "You are a document classification agent. Classify the document into one of these categories: Legal Notice, Ownership Documents, Contracts & Agreements, Financial Documents, Terms & Conditions / Privacy Policies, Intellectual Property Documents, Criminal Offense Documents, Regulatory Compliance Documents, Employment Documents, Court Judgments & Legal Precedents.",
                document_text[:3000]
            ])
            category = response.text.strip()
        else:
            category = 'default'
        return jsonify({'category': category}), 200
    except Exception as e:
        app.logger.error(f"Classification error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/process', methods=['POST'])
def process_document():
    # ensure session
    sid = _ensure_session_id()

    if 'document' not in request.files or 'category' not in request.form:
        return jsonify({'error': 'Document file or category is missing'}), 400

    pdf_file = request.files['document']
    category = request.form['category']
    pdf_bytes = pdf_file.read()

    document_text = extract_text_from_pdf_bytes(pdf_bytes)
    if not document_text or not category:
        return jsonify({'error': 'Document text or category is missing'}), 400

    # store in caches
    document_cache[sid] = document_text
    pdf_cache[sid] = pdf_bytes

    metrics = CATEGORY_METRICS.get(category, [])
    metrics_prompt = ', '.join(metrics)
    prompt = f"You are an expert summarizer for {category} documents. Extract the following relevant metrics: {metrics_prompt}. Format each metric as '**Metric Name**: Value'."

    try:
        if api_key:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content([prompt, document_text[:4000]])
            summary = response.text.strip()
        else:
            summary = "Summary disabled (no GEMINI_API_KEY)"

        # return truncated document_text so frontend can show preview
        return jsonify({
            'summary': summary,
            'document_text': document_text[:200] + '...' if len(document_text) > 200 else document_text
        }), 200
    except Exception as e:
        app.logger.error(f"Processing error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/chat', methods=['POST'])
def chat():
    # ---------------- SESSION CHECK ----------------
    sid = session.get('session_id')
    if not sid:
        sid = os.urandom(16).hex()
        session['session_id'] = sid
        app.logger.debug(f"New session created: {sid}")

    # ---------------- JSON PARSING ----------------
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': 'Invalid or missing JSON'}), 400

    user_message = data.get('message')
    category = data.get('category')
    detailed_analysis = data.get('detailed_analysis', False)
    generate_draft_flag = data.get('generate_draft', False)
    draft_instructions = data.get('draft_instructions', '')

    if not user_message:
        return jsonify({'error': 'Message is required'}), 400

    # ---------------- DOCUMENT CHECK ----------------
    document_text = document_cache.get(sid)
    if not document_text:
        return jsonify({'error': 'No document found. Please process a document first.'}), 400

    # ---------------- CHAT PROCESS ----------------
    try:
        global doc_chat_context

        doc_preview = document_text[:3000]

        system_prompt = f"""
You are a legal assistant specializing in {category} documents.
Document context: {doc_preview}
Chat history: {doc_chat_context}

Answer clearly. Use **bold** for important points.
"""

        # If user wants deeper reasoning
        if detailed_analysis:
            system_prompt += "\nProvide detailed legal reasoning.\n"
        else:
            system_prompt += "\nGive a concise answer.\n"

        # ---------------- DRAFT MODE ----------------
        if generate_draft_flag:
            draft_id = generate_document_draft(
                user_message, draft_instructions, category, doc_preview
            )
            return jsonify({
                "response": "Draft created successfully. Download below:",
                "draft_id": draft_id
            }), 200

        # ---------------- GENERATE RESPONSE ----------------
        if api_key:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content([system_prompt, user_message])
            bot_response = response.text.strip()
        else:
            bot_response = "(API disabled) " + user_message[:200]

        # ---------------- APPEND CHAT HISTORY ----------------
        doc_chat_context += f"\nUser: {user_message}\nBot: {bot_response}\n"

        # ---------------- RETURN TO FRONTEND ----------------
        return jsonify({
            "response": bot_response
        }), 200

    except Exception as e:
        app.logger.error(f"Chat error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500



@app.route('/general_chat', methods=['POST'])
def general_chat_api():
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({'error': 'Invalid or missing JSON'}), 400

    user_message = data.get('message')
    detailed_analysis = data.get('detailed_analysis', False)

    if not user_message:
        return jsonify({'error': 'Message is required'}), 400

    try:
        if detailed_analysis:
            global general_context   # <-- FIX HERE (must be first)

            response, reasoning = get_answer(user_message, general_context)

            general_context += f"\n User: {user_message}\nSenior Lawyer: {response}\n"

            return jsonify({'response': response, 'reasoning': reasoning}), 200

        else:
            general_context
            system_prompt = (
                f"You are a knowledgeable legal assistant providing general legal information.\n"
                f"You are not a lawyer and should clarify that your responses do not constitute legal advice.\n"
                f"Context: {general_context}\n"
                "Provide concise, clear answers with **bold** important points."
            )
            if api_key:
                model = genai.GenerativeModel("gemini-2.0-flash")
                response = model.generate_content([system_prompt, user_message])
                bot_response = response.text.strip()
            else:
                bot_response = "(API disabled) Mock response: " + (user_message[:200])
           
            general_context += f"\nUser: {user_message}\nBot: {bot_response}\n"
            return jsonify({'response': bot_response, 'reasoning': []}), 200
    except Exception as e:
        app.logger.error(f"General chat error: {e}")
        return jsonify({'error': str(e)}), 500


# Draft generation and formatting (same as before)

def generate_document_draft(message, instructions, category, document_context):
    template_name = CATEGORY_TO_TEMPLATE.get(category, CATEGORY_TO_TEMPLATE['default'])
    template = DRAFT_TEMPLATES[template_name]

    prompt = f"""You are a professional legal document drafter. Create a formal response document based on the instructions:
Instructions: {instructions}
Category: {category}
Context (truncated): {document_context[:1500]}
Format: header, date, salutation, body, closing, signature as per professional standards.
"""

    try:
        if api_key:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content([prompt, message])
            draft_content = response.text.strip()
        else:
            draft_content = "(API disabled) Draft content placeholder."

        doc = create_formatted_document(draft_content, template)
        draft_id = str(uuid.uuid4())
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        draft_cache[draft_id] = {
            'path': temp_file.name,
            'filename': f"Legal_Draft_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
        }
        return draft_id
    except Exception as e:
        app.logger.error(f"Draft generation error: {e}")
        raise


def create_formatted_document(content, template):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(template['margins']['top'])
        section.bottom_margin = Inches(template['margins']['bottom'])
        section.left_margin = Inches(template['margins']['left'])
        section.right_margin = Inches(template['margins']['right'])

    lines = content.split('\n')
    current_section = 'header'

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        if current_section == 'header':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template['header_format']['align'] == 'center' else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['header_format']['font']
            run.font.size = Pt(template['header_format']['size'])
            run.font.bold = template['header_format']['bold']
            current_section = 'body'
        elif current_section == 'body':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['body_format']['font']
            run.font.size = Pt(template['body']['size'] if 'body' in template else template['body_format']['size'])
        elif current_section == 'signature':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.name = template['signature_format']['font']
            run.font.size = Pt(template['signature_format']['size'])

    return doc


@app.route('/download-draft/<draft_id>', methods=['GET'])
def download_draft(draft_id):
    if draft_id not in draft_cache:
        return jsonify({'error': 'Draft not found'}), 404
    draft_info = draft_cache[draft_id]
    try:
        return send_file(
            draft_info['path'],
            as_attachment=True,
            download_name=draft_info['filename'],
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        app.logger.error(f"Draft download error: {e}")
        return jsonify({'error': 'Error downloading draft'}), 500


@app.route('/view-document', methods=['GET'])
def view_document():
    sid = session.get('session_id')
    if not sid or sid not in pdf_cache:
        return jsonify({'error': 'No document found'}), 404
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.write(pdf_cache[sid])
    temp_file.close()
    return send_file(temp_file.name, mimetype='application/pdf', as_attachment=False)


# Run app
if __name__ == '__main__':
    app.run(debug=True)
